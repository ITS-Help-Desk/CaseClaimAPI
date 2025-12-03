from rest_framework import status
from rest_framework.response import Response
from rest_framework.decorators import api_view

from rest_framework.decorators import authentication_classes, permission_classes
from rest_framework.authentication import SessionAuthentication, TokenAuthentication
from rest_framework.permissions import IsAuthenticated

from django.contrib.auth.models import User, Group
from django.db.models import Count, Q
from django.http import HttpResponse
from datetime import datetime
import statistics
import zipfile
import os
import tempfile

from user.decorators import role_required

from .models import Evaluation
from .serializers import EvaluationSerializer
from reviewedclaim.models import ReviewedClaim

# Try to import python-docx for Word document generation
try:
    from docx import Document
    from docx.shared import Pt
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


@api_view(['GET'])
def get_routes(request):
    routes = [
        {
            'Endpoint': '/create/',
            'method': 'POST',
            'body': {
                'tech': 'user_id',
                'period_start': 'YYYY-MM-DD',
                'period_end': 'YYYY-MM-DD',
                'strengths': 'text',
                'areas_for_improvement': 'text',
                'additional_comments': 'text',
                'overall_rating': '1-5'
            },
            'description': 'Create a new evaluation for a tech.'
        },
        {
            'Endpoint': '/list/',
            'method': 'GET',
            'body': None,
            'description': 'List all evaluations.'
        },
        {
            'Endpoint': '/user/<int:user_id>/',
            'method': 'GET',
            'body': None,
            'description': 'Get all evaluations for a specific user.'
        },
        {
            'Endpoint': '/detail/<int:pk>/',
            'method': 'GET',
            'body': None,
            'description': 'Get details of a specific evaluation.'
        },
        {
            'Endpoint': '/update/<int:pk>/',
            'method': 'PUT',
            'body': 'Same as create',
            'description': 'Update an existing evaluation.'
        },
        {
            'Endpoint': '/delete/<int:pk>/',
            'method': 'DELETE',
            'body': None,
            'description': 'Delete an evaluation.'
        },
        {
            'Endpoint': '/generate/<int:user_id>/',
            'method': 'GET',
            'query_params': {'start_date': 'YYYY-MM-DD', 'end_date': 'YYYY-MM-DD'},
            'description': 'Generate evaluation data/metrics for a user (does not save, just calculates).'
        },
    ]
    return Response(routes)


@api_view(['POST'])
@authentication_classes([SessionAuthentication, TokenAuthentication])
@permission_classes([IsAuthenticated])
@role_required('Lead')  # Lead and above only
def create_evaluation(request):
    """
    Create a new evaluation for a tech.
    
    The evaluator is automatically set to the current user.
    """
    serializer = EvaluationSerializer(data=request.data)
    
    if serializer.is_valid():
        # Set the evaluator to the current user
        serializer.save(evaluator=request.user)
        return Response(serializer.data, status=status.HTTP_201_CREATED)
    
    return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


@api_view(['GET'])
@authentication_classes([SessionAuthentication, TokenAuthentication])
@permission_classes([IsAuthenticated])
@role_required('Lead')  # Lead and above only
def list_evaluations(request):
    """
    List all evaluations.
    
    Supports filtering by query params:
        - tech_id: Filter by tech user ID
        - evaluator_id: Filter by evaluator user ID
    """
    evaluations = Evaluation.objects.all()
    
    # Apply filters
    tech_id = request.query_params.get('tech_id')
    evaluator_id = request.query_params.get('evaluator_id')
    
    if tech_id:
        evaluations = evaluations.filter(tech_id=tech_id)
    if evaluator_id:
        evaluations = evaluations.filter(evaluator_id=evaluator_id)
    
    serializer = EvaluationSerializer(evaluations, many=True)
    return Response(serializer.data, status=status.HTTP_200_OK)


@api_view(['GET'])
@authentication_classes([SessionAuthentication, TokenAuthentication])
@permission_classes([IsAuthenticated])
@role_required('Lead')  # Lead and above only
def get_user_evaluations(request, user_id):
    """
    Get all evaluations for a specific user (as tech).
    """
    try:
        user = User.objects.get(pk=user_id)
    except User.DoesNotExist:
        return Response({'error': 'User not found'}, status=status.HTTP_404_NOT_FOUND)
    
    evaluations = Evaluation.objects.filter(tech=user)
    serializer = EvaluationSerializer(evaluations, many=True)
    
    return Response({
        'user': {
            'id': user.id,
            'username': user.username,
            'name': f"{user.first_name} {user.last_name}".strip() or user.username
        },
        'evaluations': serializer.data,
        'total_count': evaluations.count()
    }, status=status.HTTP_200_OK)


@api_view(['GET'])
@authentication_classes([SessionAuthentication, TokenAuthentication])
@permission_classes([IsAuthenticated])
@role_required('Lead')  # Lead and above only
def get_evaluation_detail(request, pk):
    """
    Get details of a specific evaluation.
    """
    try:
        evaluation = Evaluation.objects.get(pk=pk)
    except Evaluation.DoesNotExist:
        return Response({'error': 'Evaluation not found'}, status=status.HTTP_404_NOT_FOUND)
    
    serializer = EvaluationSerializer(evaluation)
    return Response(serializer.data, status=status.HTTP_200_OK)


@api_view(['PUT', 'PATCH'])
@authentication_classes([SessionAuthentication, TokenAuthentication])
@permission_classes([IsAuthenticated])
@role_required('Lead')  # Lead and above only
def update_evaluation(request, pk):
    """
    Update an existing evaluation.
    """
    try:
        evaluation = Evaluation.objects.get(pk=pk)
    except Evaluation.DoesNotExist:
        return Response({'error': 'Evaluation not found'}, status=status.HTTP_404_NOT_FOUND)
    
    # Only the original evaluator or a Manager can update
    from user.decorators import get_user_highest_role_level, ROLE_HIERARCHY
    user_level = get_user_highest_role_level(request.user)
    manager_level = ROLE_HIERARCHY.get('Manager', 0)
    
    if evaluation.evaluator != request.user and user_level < manager_level:
        return Response(
            {'error': 'You can only update your own evaluations.'},
            status=status.HTTP_403_FORBIDDEN
        )
    
    partial = request.method == 'PATCH'
    serializer = EvaluationSerializer(evaluation, data=request.data, partial=partial)
    
    if serializer.is_valid():
        serializer.save()
        return Response(serializer.data, status=status.HTTP_200_OK)
    
    return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


@api_view(['DELETE'])
@authentication_classes([SessionAuthentication, TokenAuthentication])
@permission_classes([IsAuthenticated])
@role_required('Lead')  # Lead and above only
def delete_evaluation(request, pk):
    """
    Delete an evaluation.
    """
    try:
        evaluation = Evaluation.objects.get(pk=pk)
    except Evaluation.DoesNotExist:
        return Response({'error': 'Evaluation not found'}, status=status.HTTP_404_NOT_FOUND)
    
    # Only the original evaluator or a Manager can delete
    from user.decorators import get_user_highest_role_level, ROLE_HIERARCHY
    user_level = get_user_highest_role_level(request.user)
    manager_level = ROLE_HIERARCHY.get('Manager', 0)
    
    if evaluation.evaluator != request.user and user_level < manager_level:
        return Response(
            {'error': 'You can only delete your own evaluations.'},
            status=status.HTTP_403_FORBIDDEN
        )
    
    evaluation.delete()
    return Response({'message': 'Evaluation deleted successfully'}, status=status.HTTP_204_NO_CONTENT)


@api_view(['GET'])
@authentication_classes([SessionAuthentication, TokenAuthentication])
@permission_classes([IsAuthenticated])
@role_required('Lead')  # Lead and above only
def generate_evaluation_data(request, user_id):
    """
    Generate evaluation metrics/data for a user based on their case history.
    
    This does NOT create an evaluation - it just calculates the metrics
    that can be used to fill in an evaluation form.
    
    Query params:
        - start_date: YYYY-MM-DD
        - end_date: YYYY-MM-DD
    """
    try:
        user = User.objects.get(pk=user_id)
    except User.DoesNotExist:
        return Response({'error': 'User not found'}, status=status.HTTP_404_NOT_FOUND)
    
    # Get date range from query params
    start_date_str = request.query_params.get('start_date')
    end_date_str = request.query_params.get('end_date')
    
    if not start_date_str or not end_date_str:
        return Response({
            'error': 'Both start_date and end_date are required (YYYY-MM-DD format)'
        }, status=status.HTTP_400_BAD_REQUEST)
    
    try:
        start_date = datetime.strptime(start_date_str, '%Y-%m-%d')
        end_date = datetime.strptime(end_date_str, '%Y-%m-%d')
        end_date = end_date.replace(hour=23, minute=59, second=59)
    except ValueError:
        return Response({
            'error': 'Invalid date format. Use YYYY-MM-DD'
        }, status=status.HTTP_400_BAD_REQUEST)
    
    # Get all reviewed claims for this tech in the date range
    reviewed_claims = ReviewedClaim.objects.filter(
        tech_id=user,
        review_time__gte=start_date,
        review_time__lte=end_date
    )
    
    total_cases = reviewed_claims.count()
    
    # Count by status
    kudos_count = reviewed_claims.filter(status='kudos').count()
    checked_count = reviewed_claims.filter(status='checked').count()
    done_count = reviewed_claims.filter(status='done').count()
    
    ping_statuses = ['pingedlow', 'pingedmed', 'pingedhigh', 'acknowledged', 'resolved']
    ping_count = reviewed_claims.filter(status__in=ping_statuses).count()
    
    # Calculate quality score (simple formula: (positive / total) * 100)
    positive_reviews = kudos_count + checked_count + done_count
    quality_score = round((positive_reviews / total_cases * 100), 2) if total_cases > 0 else None
    
    # Calculate ping rate
    ping_rate = round((ping_count / total_cases * 100), 2) if total_cases > 0 else 0
    
    evaluation_data = {
        'user': {
            'id': user.id,
            'username': user.username,
            'name': f"{user.first_name} {user.last_name}".strip() or user.username
        },
        'period': {
            'start_date': start_date_str,
            'end_date': end_date_str
        },
        'metrics': {
            'cases_reviewed': total_cases,
            'quality_score': quality_score,
            'ping_count': ping_count,
            'ping_rate': ping_rate,
            'kudos_count': kudos_count,
        },
        'breakdown': {
            'kudos': kudos_count,
            'checked': checked_count,
            'done': done_count,
            'pinged_low': reviewed_claims.filter(status='pingedlow').count(),
            'pinged_med': reviewed_claims.filter(status='pingedmed').count(),
            'pinged_high': reviewed_claims.filter(status='pingedhigh').count(),
            'acknowledged': reviewed_claims.filter(status='acknowledged').count(),
            'resolved': reviewed_claims.filter(status='resolved').count(),
        },
        'suggested_rating': _calculate_suggested_rating(quality_score, ping_rate),
    }
    
    return Response(evaluation_data, status=status.HTTP_200_OK)


def _calculate_suggested_rating(quality_score, ping_rate):
    """
    Calculate a suggested rating based on metrics.
    
    This is just a suggestion - the lead can override it.
    """
    if quality_score is None:
        return None
    
    if quality_score >= 95 and ping_rate < 5:
        return 5  # Outstanding
    elif quality_score >= 85 and ping_rate < 10:
        return 4  # Exceeds Expectations
    elif quality_score >= 70 and ping_rate < 20:
        return 3  # Meets Expectations
    elif quality_score >= 50:
        return 2  # Below Expectations
    else:
        return 1  # Needs Improvement


# =====================
# GENEVAL - Auto-generate evaluations
# =====================

def _month_number_to_name(month: int) -> str:
    """Convert month number to name."""
    months = ['', 'January', 'February', 'March', 'April', 'May', 'June',
              'July', 'August', 'September', 'October', 'November', 'December']
    return months[month] if 1 <= month <= 12 else ''


def _is_user_lead(user):
    """Check if user has Lead or higher role."""
    lead_groups = ['Lead', 'Phone Analyst', 'Manager']
    return user.groups.filter(name__in=lead_groups).exists()


def _get_eval_data(month: int, year: int):
    """
    Get all evaluation data for a given month/year.
    Returns data organized by tech.
    """
    # Calculate date range
    start_date = datetime(year, month, 1)
    if month == 12:
        end_date = datetime(year + 1, 1, 1)
    else:
        end_date = datetime(year, month + 1, 1)
    
    # Get all reviewed claims for the month
    all_claims = ReviewedClaim.objects.filter(
        review_time__gte=start_date,
        review_time__lt=end_date
    ).exclude(status='done')  # Exclude 'done' status like the bot does
    
    total_hd_cases = 0
    tech_data = {}
    
    for claim in all_claims:
        tech = claim.tech_id
        
        # Skip if tech is a lead
        if _is_user_lead(tech):
            continue
        
        total_hd_cases += 1
        
        # Initialize tech data if not exists
        if tech.id not in tech_data:
            tech_data[tech.id] = {
                'user': tech,
                'checked_count': 0,
                'pinged_cases': [],
                'kudos_cases': [],
            }
        
        # Categorize by status
        if claim.status == 'checked':
            tech_data[tech.id]['checked_count'] += 1
        elif claim.status in ['pingedlow', 'pingedmed', 'pingedhigh', 'acknowledged', 'resolved']:
            tech_data[tech.id]['pinged_cases'].append(claim.casenum)
        elif claim.status == 'kudos':
            tech_data[tech.id]['kudos_cases'].append(claim.casenum)
    
    return total_hd_cases, tech_data


def _organize_data_for_word(total_case_count: int, tech_data: dict):
    """
    Organize tech data for Word document generation.
    Calculate team statistics.
    """
    if total_case_count == 0:
        return 0, 0, 0, 0, {}
    
    organized_data = {}
    claim_median_list = []
    top_claim = 0
    ping_median_list = []
    
    for tech_id, data in tech_data.items():
        claim_count = len(data['pinged_cases']) + len(data['kudos_cases']) + data['checked_count']
        
        if claim_count == 0:
            continue
        
        if claim_count > top_claim:
            top_claim = claim_count
        
        claim_percent = claim_count / total_case_count
        ping_percent = len(data['pinged_cases']) / claim_count
        
        claim_median_list.append(claim_percent)
        ping_median_list.append(ping_percent)
        
        organized_data[tech_id] = {
            'user': data['user'],
            'user_pinged_count': len(data['pinged_cases']),
            'user_total_count': claim_count,
            'claim_percent': claim_percent,
            'pings': data['pinged_cases'],
            'kudos': data['kudos_cases'],
            'ping_percent': ping_percent
        }
    
    median_claim = statistics.median(claim_median_list) if claim_median_list else 0
    median_ping = statistics.median(ping_median_list) if ping_median_list else 0
    top_claim_percent = top_claim / total_case_count if total_case_count > 0 else 0
    
    return total_case_count, median_claim, median_ping, top_claim_percent, organized_data


def _create_word_fields(username, date, hd_total_claims, median_claim, median_ping_percent, top_claim_percent, data):
    """Create field mappings for Word document template."""
    kudos = data.get('kudos', [])
    pings = data.get('pings', [])
    
    checkmark = " " + "\u2705"
    
    fields = {
        18: username,
        20: 0,
        31: date.strftime("%m/%d/%Y"),
        37: kudos[0] + checkmark if len(kudos) > 0 else "",
        38: kudos[1] + checkmark if len(kudos) > 1 else "",
        39: kudos[2] + checkmark if len(kudos) > 2 else "",
        40: kudos[3] + checkmark if len(kudos) > 3 else "",
        41: kudos[4] + checkmark if len(kudos) > 4 else "",
        52: pings[0] if len(pings) > 0 else "",
        53: pings[1] if len(pings) > 1 else "",
        54: pings[2] if len(pings) > 2 else "",
        55: pings[3] if len(pings) > 3 else "",
        56: pings[4] if len(pings) > 4 else "",
    }
    
    template_name = "templates/template1row.docx"
    offset = 0
    
    if len(pings) > 5:
        template_name = "templates/template2row.docx"
        offset = 5
        fields[57] = pings[5] if len(pings) > 5 else ""
        fields[58] = pings[6] if len(pings) > 6 else ""
        fields[59] = pings[7] if len(pings) > 7 else ""
        fields[60] = pings[8] if len(pings) > 8 else ""
        fields[61] = pings[9] if len(pings) > 9 else ""
    
    if len(pings) > 10:
        template_name = "templates/template3row.docx"
        offset = 10
        fields[62] = pings[10] if len(pings) > 10 else ""
        fields[63] = pings[11] if len(pings) > 11 else ""
        fields[64] = pings[12] if len(pings) > 12 else ""
        fields[65] = pings[13] if len(pings) > 13 else ""
        fields[66] = pings[14] if len(pings) > 14 else ""
    
    if len(pings) > 15:
        template_name = "templates/template4row.docx"
        offset = 15
        fields[67] = pings[15] if len(pings) > 15 else ""
        fields[68] = pings[16] if len(pings) > 16 else ""
        fields[69] = pings[17] if len(pings) > 17 else ""
        fields[70] = pings[18] if len(pings) > 18 else ""
        fields[71] = pings[19] if len(pings) > 19 else ""
    
    if len(pings) > 20:
        template_name = "templates/template5row.docx"
        offset = 20
        fields[72] = pings[20] if len(pings) > 20 else ""
        fields[73] = pings[21] if len(pings) > 21 else ""
        fields[74] = pings[22] if len(pings) > 22 else ""
        fields[75] = pings[23] if len(pings) > 23 else ""
        fields[76] = pings[24] if len(pings) > 24 else ""
    
    # Top stats label
    fields[62 + offset] = f"{data['user_pinged_count']} / {data['user_total_count']} / {hd_total_claims}"
    
    # User case stats
    fields[65 + offset] = f"Individual: {data['claim_percent']:.2%}"
    fields[66 + offset] = f"Team Median: {median_claim:.2%}"
    fields[67 + offset] = f"Team Top: {top_claim_percent:.2%}"
    
    # Ping stats
    fields[70 + offset] = f"Individual: {data['ping_percent']:.2%}"
    fields[71 + offset] = f"Team Median: {median_ping_percent:.2%}"
    
    return fields, template_name


def _create_document(template_path: str, save_path: str, fields: dict):
    """Create a Word document from template with filled fields."""
    document = Document(template_path)
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    
    i = 1
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if i in fields:
                        paragraph.text = str(fields[i])
                        paragraph.style = document.styles['Normal']
                    i += 1
    
    document.save(save_path)


@api_view(['GET'])
@authentication_classes([SessionAuthentication, TokenAuthentication])
@permission_classes([IsAuthenticated])
@role_required('Lead')
def geneval(request):
    """
    Generate evaluation documents for all techs for a given month/year.
    
    Query params:
        - month: 1-12
        - year: YYYY
    
    Returns a ZIP file containing Word documents for each tech.
    """
    if not DOCX_AVAILABLE:
        return Response({
            'error': 'python-docx library is not installed. Please install it with: pip install python-docx'
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    
    # Get month and year from query params
    try:
        month = int(request.query_params.get('month'))
        year = int(request.query_params.get('year'))
    except (TypeError, ValueError):
        return Response({
            'error': 'Both month (1-12) and year (YYYY) are required as query parameters'
        }, status=status.HTTP_400_BAD_REQUEST)
    
    if not 1 <= month <= 12:
        return Response({'error': 'Month must be between 1 and 12'}, status=status.HTTP_400_BAD_REQUEST)
    
    if not 2000 <= year <= 2100:
        return Response({'error': 'Year must be between 2000 and 2100'}, status=status.HTTP_400_BAD_REQUEST)
    
    # Get evaluation data
    total_hd_cases, tech_data = _get_eval_data(month, year)
    
    if total_hd_cases == 0:
        return Response({
            'error': f'No cases found for {_month_number_to_name(month)} {year}'
        }, status=status.HTTP_404_NOT_FOUND)
    
    # Organize data for Word documents
    hd_total_claims, median_claim, median_ping_percent, top_claim_percent, organized_data = \
        _organize_data_for_word(total_hd_cases, tech_data)
    
    if not organized_data:
        return Response({
            'error': f'No tech data found for {_month_number_to_name(month)} {year}'
        }, status=status.HTTP_404_NOT_FOUND)
    
    # Create temp directory for files
    with tempfile.TemporaryDirectory() as temp_dir:
        now = datetime.now()
        filenames = []
        
        # Get base directory for templates
        base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        
        for tech_id, data in organized_data.items():
            user = data['user']
            full_name = f"{user.first_name} {user.last_name}".strip() or user.username
            
            fields, template_name = _create_word_fields(
                full_name,
                now,
                hd_total_claims,
                median_claim,
                median_ping_percent,
                top_claim_percent,
                data
            )
            
            # Create document
            template_path = os.path.join(base_dir, template_name)
            last_name = full_name.split(' ')[-1] if ' ' in full_name else full_name
            doc_filename = f"{last_name} {_month_number_to_name(month).upper()} {year}.docx"
            doc_path = os.path.join(temp_dir, doc_filename)
            
            try:
                _create_document(template_path, doc_path, fields)
                filenames.append((doc_path, doc_filename))
            except Exception as e:
                print(f"Error creating document for {full_name}: {e}")
                continue
        
        if not filenames:
            return Response({
                'error': 'Failed to generate any evaluation documents'
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        
        # Create ZIP file
        zip_filename = f"evaluations_{month}_{year}.zip"
        zip_path = os.path.join(temp_dir, zip_filename)
        
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for file_path, filename in filenames:
                zipf.write(file_path, filename)
        
        # Read ZIP and return as response
        with open(zip_path, 'rb') as f:
            zip_data = f.read()
        
        response = HttpResponse(zip_data, content_type='application/zip')
        response['Content-Disposition'] = f'attachment; filename="{zip_filename}"'
        return response
